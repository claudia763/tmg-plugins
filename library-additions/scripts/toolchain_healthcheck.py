#!/usr/bin/env python3
"""
Toolchain health check for the TMG email-cowork agent server.

WHAT IT DOES
    Verifies that everything the multifamily-brokerage skills depend on is
    actually installed AND functional on this machine -- not just importable.
    It writes real files (xlsx/docx/pdf) and launches headless Chromium, because
    the common failure mode is a package that imports fine but a binary that
    cannot start (e.g. Chromium missing system .so files on a no-root host).

WHEN TO USE
    - First thing on a new/rebuilt server, or when a job fails with an
      ImportError / "error while loading shared libraries" / browser launch error.
    - After any toolchain change, to confirm nothing regressed.

INPUTS
    --skills-dir PATH   Optional. Path to multifamily-brokerage/skills so the
                        T-12 / rent-roll templates and mapping table are checked
                        too. Auto-detected if omitted.

OUTPUTS
    Prints one PASS/FAIL line per check plus a summary.
    Exit code 0 = all checks passed, 1 = one or more FAILed.
    Scratch files are written to a temp dir and deleted.

NOTES
    This host (ubuntu-8gb-nbg1-2) has no sudo. If a check fails, see
    instructions/no-root-toolchain-setup.md in library-additions for the
    rootless fix (pip --user, npm user prefix, .deb extraction + patchelf).
"""

import argparse
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

RESULTS = []


def record(name, ok, detail=""):
    RESULTS.append((name, ok, detail))
    print(f"  [{'PASS' if ok else 'FAIL'}] {name}" + (f" -- {detail}" if detail else ""))
    return ok


def check_python_imports():
    """Each package the brokerage skills import at runtime."""
    wanted = {
        "openpyxl": "rent roll / T-12 workbook read+write",
        "pandas": "tabular processing",
        "numpy": "numerics",
        "fitz": "PyMuPDF -- PDF read + deck verification",
        "pdfplumber": "PDF table extraction (rent rolls, T-12s)",
        "docx": "python-docx -- Word deliverables",
        "PIL": "Pillow -- photo prep for decks",
        "rapidfuzz": "fuzzy GL-account matching (t12_mappings)",
    }
    for mod, why in wanted.items():
        try:
            m = __import__(mod)
            record(f"import {mod}", True, f"{getattr(m, '__version__', '')} ({why})".strip())
        except Exception as e:
            record(f"import {mod}", False, f"{e} ({why})")


def check_file_writes(tmp):
    """Prove we can actually produce each deliverable file type."""
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"], ws["A2"] = 1, "=A1+1"
        wb.save(tmp / "t.xlsx")
        openpyxl.load_workbook(tmp / "t.xlsx")
        record("write+reopen .xlsx", True)
    except Exception as e:
        record("write+reopen .xlsx", False, str(e))

    try:
        import docx
        d = docx.Document()
        d.add_heading("h", 0)
        d.add_paragraph("p")
        d.save(tmp / "t.docx")
        record("write .docx", True)
    except Exception as e:
        record("write .docx", False, str(e))

    try:
        import fitz
        doc = fitz.open()
        doc.new_page().insert_text((72, 72), "probe")
        doc.save(tmp / "t.pdf")
        doc.close()
        with fitz.open(tmp / "t.pdf") as r:
            got = r[0].get_text().strip()
        record("write+read .pdf", got == "probe", f"round-tripped text={got!r}")
    except Exception as e:
        record("write+read .pdf", False, str(e))


def check_node():
    node = shutil.which("node")
    if not node:
        record("node runtime", False, "node not on PATH")
        return False
    v = subprocess.run([node, "-v"], capture_output=True, text=True).stdout.strip()
    record("node runtime", True, v)
    for mod in ("playwright", "docx"):
        r = subprocess.run([node, "-e", f"require('{mod}')"], capture_output=True, text=True)
        record(f"node require('{mod}')", r.returncode == 0,
               "" if r.returncode == 0 else r.stderr.strip().splitlines()[0][:120])
    return True


def check_chromium(tmp):
    """The check that matters most: does headless Chromium actually START?

    Library skills (bov-deck, comp maps) render through Playwright. A missing
    system .so shows up only at launch, never at import.
    """
    node = shutil.which("node")
    if not node:
        return record("chromium headless render", False, "no node")
    script = tmp / "probe.cjs"
    script.write_text(
        "const {chromium}=require('playwright');(async()=>{"
        "const b=await chromium.launch();const p=await b.newPage({viewport:{width:400,height:200}});"
        "await p.setContent('<h1>probe</h1>');"
        f"await p.screenshot({{path:{str(tmp / 'probe.png')!r}}});"
        f"await p.pdf({{path:{str(tmp / 'probe_deck.pdf')!r},width:'13.333in',height:'7.5in'}});"
        "await b.close();console.log('ok');})()"
        ".catch(e=>{console.error(e.message);process.exit(1)});"
    )
    r = subprocess.run([node, str(script)], capture_output=True, text=True, timeout=180)
    if r.returncode != 0:
        first = (r.stderr.strip().splitlines() or ["launch failed"])[0]
        return record("chromium headless render", False, first[:160])
    png_ok = (tmp / "probe.png").stat().st_size > 500
    record("chromium headless render (png)", png_ok)
    try:
        import fitz
        with fitz.open(tmp / "probe_deck.pdf") as d:
            w, h = d[0].rect.width, d[0].rect.height
        # 13.333in x 7.5in at 72pt/in = 960 x 540 -- the deck page geometry
        record("chromium PDF render (960x540pt deck page)", round(w) == 960 and round(h) == 540,
               f"{w:.0f}x{h:.0f}pt")
    except Exception as e:
        record("chromium PDF render", False, str(e))
    return True


def autodetect_skills_dir():
    for c in (
        Path("/home/claudia/email-cowork-server/resources/tmg-plugins/multifamily-brokerage/skills"),
        Path(__file__).resolve().parent.parent / "multifamily-brokerage" / "skills",
    ):
        if c.is_dir():
            return c
    return None


def check_library(skills_dir):
    if not skills_dir or not Path(skills_dir).is_dir():
        record("library skills dir", False, f"not found: {skills_dir}")
        return
    skills_dir = Path(skills_dir)
    record("library skills dir", True, str(skills_dir))
    try:
        import warnings
        import openpyxl
        for name in ("t12_processor_template.xlsx", "rentroll_template.xlsx"):
            p = skills_dir / name
            if not p.exists():
                record(f"template {name}", False, "missing")
                continue
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                wb = openpyxl.load_workbook(p)
            record(f"template {name}", True, f"{len(wb.sheetnames)} tabs")
    except Exception as e:
        record("templates", False, str(e))

    m = skills_dir / "t12_mappings.csv"
    if m.exists():
        try:
            import csv
            rows = sum(1 for _ in csv.reader(m.open(encoding="utf-8-sig"))) - 1
            record("t12_mappings.csv", rows > 0, f"{rows} GL mapping rows")
        except Exception as e:
            record("t12_mappings.csv", False, str(e))
    else:
        record("t12_mappings.csv", False, "missing")

    for s in ("process_t12.py", "process_rent_roll.py"):
        p = skills_dir / s
        if not p.exists():
            record(f"compile {s}", False, "missing")
            continue
        r = subprocess.run([sys.executable, "-m", "py_compile", str(p)], capture_output=True, text=True)
        record(f"compile {s}", r.returncode == 0, r.stderr.strip()[:120])


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--skills-dir", default=None, help="path to multifamily-brokerage/skills")
    args = ap.parse_args()

    print(f"TMG toolchain health check -- host={os.uname().nodename} python={sys.version.split()[0]}\n")
    print("Python packages:")
    check_python_imports()
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        print("\nDeliverable file types:")
        check_file_writes(tmp)
        print("\nNode / rendering:")
        check_node()
        check_chromium(tmp)
    print("\nLibrary assets:")
    check_library(args.skills_dir or autodetect_skills_dir())

    failed = [n for n, ok, _ in RESULTS if not ok]
    print(f"\n{'-' * 60}")
    print(f"{len(RESULTS) - len(failed)}/{len(RESULTS)} checks passed")
    if failed:
        print("FAILED: " + ", ".join(failed))
        print("See instructions/no-root-toolchain-setup.md for the rootless fix.")
    else:
        print("All systems go.")
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
